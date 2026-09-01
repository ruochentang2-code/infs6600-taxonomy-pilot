"""Generate the Week 4 eight-person work allocation meeting handout."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("output/docx/CS-44_Week4_Eight-Person_Work_Allocation.docx")

NAVY = "1F4D78"
BLUE = "2E74B5"
BLUE_DARK = "17365D"
BLUE_LIGHT = "DCE6F1"
PALE_BLUE = "EEF4F8"
GOLD = "D6A84B"
PALE_GOLD = "FFF4D6"
GREY = "F2F4F7"
MID_GREY = "D0D7DE"
TEXT = "243B53"
MUTED = "52606D"
WHITE = "FFFFFF"
GREEN = "2E7D32"


MEMBERS = [
    {
        "number": "01",
        "name": "Member 1",
        "stream": "Coordination, decision traceability and scope control",
        "stream_zh": "项目协调、决策追踪与范围控制",
        "completed": [
            "Converted the meeting, email, PDF comments and scope-document changes into one acceptance checklist.",
            "Kept Week 4 limited to the agreed INFS6600 update and its necessary supporting work.",
            "Coordinated the eight workstreams, hand-offs and shared meeting narrative.",
        ],
        "outputs": [
            "Meeting decision record and Week 4 implementation record",
            "Week 4 delivery checklist and scope guard",
            "Integrated meeting handout and speaking order",
        ],
        "completed_zh": [
            "将会议、邮件、PDF 批注及范围文档修订整合为统一验收清单。",
            "将第四周工作限定在已确认的 INFS6600 更新及必要配套工作内。",
            "协调八个工作流、任务交接与统一会议汇报口径。",
        ],
        "outputs_zh": [
            "会议决策记录与第四周实施记录",
            "第四周交付清单与范围控制说明",
            "整合后的会议材料与汇报顺序",
        ],
        "line": "I coordinated the Week 4 v2 update, mapped every agreed change to an acceptance check, and kept the team within the INFS6600-only Week 4 scope.",
        "line_zh": "我负责协调第四周 v2 更新，将每项已确认修改对应到验收检查，并确保团队严格保持在 INFS6600 第四周范围内。",
    },
    {
        "number": "02",
        "name": "Member 2",
        "stream": "Taxonomy definitions and versioned configuration",
        "stream_zh": "分类体系定义与版本化配置",
        "completed": [
            "Configured all eight supplied taxonomy categories for the pilot.",
            "Separated Simulation from Case-Based Learning and clarified their definitions.",
            "Added overlap notes, review guidance and provisional-status labels to the taxonomy.",
        ],
        "outputs": [
            "config/taxonomy_v2.json",
            "Eight-category standards in the detailed algorithm report",
            "Notes for open configuration points",
        ],
        "completed_zh": [
            "为试点配置全部八个既定分类。",
            "将模拟学习与案例式学习拆分，并明确两者定义。",
            "在分类体系中补充类别重叠说明、人工复核指引和临时状态标记。",
        ],
        "outputs_zh": [
            "config/taxonomy_v2.json",
            "详细算法报告中的八分类标准",
            "待客户确认事项的配置说明",
        ],
        "line": "I updated the supplied taxonomy into a versioned eight-category configuration and made Simulation and Case-Based Learning independent categories with clear overlap and review guidance.",
        "line_zh": "我将分类体系更新为版本化的八分类配置，并把模拟学习与案例式学习设为独立类别，同时补充清晰的重叠规则和复核指引。",
    },
    {
        "number": "03",
        "name": "Member 3",
        "stream": "Multi-label, overlap and review policy",
        "stream_zh": "多标签、类别重叠与人工复核政策",
        "completed": [
            "Defined categories as non-mutually-exclusive and removed any single-winner assumption.",
            "Retained an evidence item once per category so genuine overlap remains visible.",
            "Moved administrative “Case studies” labels to review rather than treating them as automatic positive evidence.",
        ],
        "outputs": [
            "Multi-label and item-category deduplication policy",
            "Case/industry overlap interpretation",
            "Manual-review rules and review-queue criteria",
        ],
        "completed_zh": [
            "明确各分类并非互斥，并移除单一胜出类别的假设。",
            "同一证据条目可在每个相关类别中计数一次，以保留真实重叠关系。",
            "将行政标签“Case studies”转入人工复核，而非自动作为正向证据。",
        ],
        "outputs_zh": [
            "多标签与“证据条目—类别”去重政策",
            "案例与行业合作重叠关系解释",
            "人工复核规则与复核队列标准",
        ],
        "line": "I formalised the multi-label policy: a unit or evidence item can support several categories, while ambiguous administrative labels stay visible in the review queue instead of becoming automatic positives.",
        "line_zh": "我正式确立了多标签政策：一个课程或证据条目可以支持多个类别，而含义不明确的行政标签会保留在复核队列中，不会自动成为正向结果。",
    },
    {
        "number": "04",
        "name": "Member 4",
        "stream": "INFS6600 source evidence and pilot rerun",
        "stream_zh": "INFS6600 来源证据与试点重跑",
        "completed": [
            "Verified the INFS6600 2026 Semester 2 public-outline snapshot and official source URL.",
            "Maintained auditable evidence units across overview, learning outcomes, assessments and weekly schedule.",
            "Reran the pilot and checked item IDs, source sections and original evidence text.",
        ],
        "outputs": [
            "Versioned INFS6600 source snapshot and course-information summary",
            "Classified-evidence and review-queue datasets",
            "Traceable evidence for the three positive categories",
        ],
        "completed_zh": [
            "核验 INFS6600 2026 年第二学期公开课程大纲快照及官方来源链接。",
            "维护课程概述、学习成果、考核任务和每周安排中的可审计证据条目。",
            "重新运行试点，并检查条目编号、来源章节和原始证据文本。",
        ],
        "outputs_zh": [
            "版本化 INFS6600 来源快照与课程信息摘要",
            "已分类证据与人工复核队列数据集",
            "三个正向类别的可追溯证据",
        ],
        "line": "I prepared and verified the INFS6600 evidence base, reran the pilot across all four outline sections, and ensured every result remained traceable to an item ID and the official source.",
        "line_zh": "我准备并核验了 INFS6600 证据库，在课程大纲的四个部分重新运行试点，并确保每项结果均可追溯到条目编号和官方来源。",
    },
    {
        "number": "05",
        "name": "Member 5",
        "stream": "Scoring method, WIL weights and confidence logic",
        "stream_zh": "评分方法、WIL 权重与置信度逻辑",
        "completed": [
            "Calibrated the four WIL weights to 4.0, 2.0, 2.0 and provisional 2.0.",
            "Kept one contribution per matched rule group and separated classified score from review score.",
            "Documented the 3.0 and 5.0 thresholds as configurable and provisional.",
        ],
        "outputs": [
            "Updated weighted phrase-scoring configuration and logic",
            "Classified, review and total matched score fields",
            "02_Classification_Algorithm_Detailed_v2.pdf",
        ],
        "completed_zh": [
            "将四项 WIL 权重校准为 4.0、2.0、2.0 和临时 2.0。",
            "每个匹配规则组仅计分一次，并将分类得分与复核得分分开。",
            "明确 3.0 和 5.0 阈值可配置且仍属临时设置。",
        ],
        "outputs_zh": [
            "更新后的加权短语评分配置与逻辑",
            "分类得分、复核得分和总匹配得分字段",
            "02_Classification_Algorithm_Detailed_v2.pdf",
        ],
        "line": "I calibrated the WIL weights, separated positive and review scoring, and kept the decision thresholds explicit and provisional rather than presenting them as validated cut-offs.",
        "line_zh": "我校准了 WIL 权重，将正向评分与复核评分分开，并明确决策阈值仍为临时设置，而非已经验证的界限。",
    },
    {
        "number": "06",
        "name": "Member 6",
        "stream": "Evidence aggregation and course-category mapping",
        "stream_zh": "证据汇总与课程—类别映射",
        "completed": [
            "Counted distinct evidence items rather than raw keyword occurrences.",
            "Aggregated positive items, review items, classified score, review score and section counts by category.",
            "Produced the complete eight-category mapping while preserving the three positive INFS6600 allocations.",
        ],
        "outputs": [
            "classification_results.json and unit_category_summary.csv",
            "Course-category mapping in Markdown and CSV",
            "Verified WIL, Case-Based and Project/Problem-Based positive result",
        ],
        "completed_zh": [
            "统计不同证据条目，而不是原始关键词出现次数。",
            "按类别汇总正向条目、复核条目、分类得分、复核得分和来源章节数量。",
            "生成完整八分类映射，同时保留 INFS6600 的三个正向类别结果。",
        ],
        "outputs_zh": [
            "classification_results.json 与 unit_category_summary.csv",
            "Markdown 和 CSV 格式的课程—类别映射",
            "经核验的 WIL、案例式学习和项目/问题式学习正向结果",
        ],
        "line": "I built the category-level aggregation and mapping outputs, using distinct evidence-item counts and separate score totals so the three INFS6600 allocations can be audited clearly.",
        "line_zh": "我构建了类别级汇总和映射输出，采用不同证据条目数量及分离的得分合计，使 INFS6600 的三个分类结果能够被清晰审计。",
    },
    {
        "number": "07",
        "name": "Member 7",
        "stream": "Visualisation and report production",
        "stream_zh": "可视化与报告制作",
        "completed": [
            "Created three INFS6600-only taxonomy figures for evidence counts, scores and source sections.",
            "Integrated total scoring, review signals and multi-label interpretation into the pilot report.",
            "Produced the updated course-category mapping report for meeting review.",
        ],
        "outputs": [
            "Three versioned PNG figures",
            "pilot_results_v2.md",
            "05_INFS6600_Course_Category_Mapping_v2.pdf",
        ],
        "completed_zh": [
            "为证据数量、类别得分和来源章节制作三张仅针对 INFS6600 的分类图。",
            "将总评分、复核信号和多标签解释整合进试点报告。",
            "制作更新后的课程—类别映射报告，供会议审阅。",
        ],
        "outputs_zh": [
            "三张版本化 PNG 图表",
            "pilot_results_v2.md",
            "05_INFS6600_Course_Category_Mapping_v2.pdf",
        ],
        "line": "I turned the updated results into the three requested pilot-scale figures and the mapping report, showing evidence counts, total scores and where the evidence appears in the unit outline.",
        "line_zh": "我将更新结果制作成三张试点规模图表和映射报告，展示证据数量、总得分以及证据在课程大纲中的分布位置。",
    },
    {
        "number": "08",
        "name": "Member 8",
        "stream": "Quality assurance, release packaging and version control",
        "stream_zh": "质量保证、发布打包与版本控制",
        "completed": [
            "Ran offline unit and Week 4 regression checks for the eight-category multi-label result.",
            "Verified generated reports, release hashes and the two final PDF deliverables.",
            "Packaged the update as a separate Week 4 v2 Git version while preserving the original main branch.",
        ],
        "outputs": [
            "Regression tests and acceptance-check results",
            "SHA-256 release manifest, README and changelog",
            "Versioned Git branch and final delivery package",
        ],
        "completed_zh": [
            "对八分类多标签结果运行离线单元测试和第四周回归检查。",
            "核验生成的报告、发布哈希和两份最终 PDF 交付物。",
            "以独立的 Week 4 v2 Git 版本打包更新，并保留原始 main 分支。",
        ],
        "outputs_zh": [
            "回归测试与验收检查结果",
            "SHA-256 发布清单、README 和更新日志",
            "版本化 Git 分支与最终交付包",
        ],
        "line": "I completed release QA and version packaging, confirmed the expected three-category result and kept this Week 4 update separate so the original repository version remains intact.",
        "line_zh": "我完成了发布质量检查和版本打包，确认预期的三个正向类别结果，并将本次第四周更新作为独立版本，确保原始仓库版本保持完整。",
    },
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = f"w:{edge}"
        node = tc_mar.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_fixed(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            cell.width = Inches(widths[idx] / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_border(cell, color=MID_GREY, size="6", **edges) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        settings = edges.get(edge, {"val": "single", "sz": size, "color": color})
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key, value in settings.items():
            element.set(qn(f"w:{key}"), str(value))


DOCUMENT_FONT = "Hiragino Sans GB"


def set_run_font(run, font_name: str = DOCUMENT_FONT) -> None:
    """Set every OOXML font slot so Word and LibreOffice both render CJK text."""
    run.font.name = font_name
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for slot in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        fonts.set(qn(slot), font_name)


def set_style_font(style, font_name: str = DOCUMENT_FONT) -> None:
    style.font.name = font_name
    fonts = style._element.get_or_add_rPr().get_or_add_rFonts()
    for slot in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        fonts.set(qn(slot), font_name)


def add_run(paragraph, text: str, *, bold=False, color=TEXT, size=None, italic=False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    set_run_font(run)
    run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)
    return run


def clear_cell(cell) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)


def add_bullet(cell, text: str, *, compact=True) -> None:
    paragraph = cell.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.first_line_indent = Inches(-0.12)
    paragraph.paragraph_format.space_after = Pt(2 if compact else 4)
    paragraph.paragraph_format.line_spacing = 1.0
    add_run(paragraph, text, size=8.6)


def add_bilingual_bullet(cell, english: str, chinese: str) -> None:
    paragraph = cell.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.first_line_indent = Inches(-0.12)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.0
    add_run(paragraph, english, size=8.15)
    add_run(paragraph, f"\n{chinese}", color=MUTED, size=8.15)


def add_field(paragraph, field: str) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = field
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr_text, fld_end])


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.32)
    section.footer_distance = Inches(0.35)

    styles = document.styles
    normal = styles["Normal"]
    set_style_font(normal)
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    title = styles["Title"]
    set_style_font(title)
    title.font.size = Pt(25)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(NAVY)
    title.paragraph_format.space_after = Pt(5)

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 15, 7),
        ("Heading 2", 13, NAVY, 11, 5),
        ("Heading 3", 11.5, BLUE_DARK, 8, 4),
    ):
        style = styles[style_name]
        set_style_font(style)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    list_style = styles["List Bullet"]
    set_style_font(list_style)
    list_style.font.size = Pt(9)
    list_style.paragraph_format.space_after = Pt(3)

    for sec in document.sections:
        header = sec.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(3)
        add_run(p, "CS-44", bold=True, color=NAVY, size=8.5)
        add_run(p, "  |  Week 4 v2 / 第四周 v2  |  Eight-Person Work Allocation / 八人分工", color=MUTED, size=8.1)
        p_pr = p._p.get_or_add_pPr()
        p_bdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "8")
        bottom.set(qn("w:color"), BLUE)
        bottom.set(qn("w:space"), "3")
        p_bdr.append(bottom)
        p_pr.append(p_bdr)

        footer = sec.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(fp, "CS-44 Project Team / 项目团队（8人）  •  1 September 2026 / 2026年9月1日  •  Page / 页 ", color=MUTED, size=7.6)
        add_field(fp, "PAGE")


def add_metadata(document: Document) -> None:
    table = document.add_table(rows=2, cols=4)
    set_table_fixed(table, [1200, 3480, 1200, 3480])
    values = [
        ("Project / 项目", "Exploring Pedagogical Innovations in Business Education\n商业教育中的教学创新探索", "Group / 小组", "CS-44"),
        ("Meeting / 会议", "Week 4 progress report / 第四周进度汇报", "Release / 版本", "week4-v2"),
    ]
    for row, values_row in zip(table.rows, values):
        prevent_row_split(row)
        for idx, value in enumerate(values_row):
            cell = row.cells[idx]
            clear_cell(cell)
            set_cell_shading(cell, GREY if idx % 2 == 0 else WHITE)
            set_cell_border(cell)
            p = cell.paragraphs[0]
            add_run(p, value, bold=idx % 2 == 0, color=NAVY if idx % 2 == 0 else TEXT, size=8.5)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_callout(document: Document, text: str, *, fill=BLUE_LIGHT, border=BLUE) -> None:
    table = document.add_table(rows=1, cols=1)
    prevent_row_split(table.rows[0])
    set_table_fixed(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, color=border, size="10")
    clear_cell(cell)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_run(p, text, bold=True, color=NAVY, size=10)


def add_page_heading(document: Document, text: str, *, level: int = 1):
    """Start a new page without inserting a standalone break paragraph."""
    heading = document.add_heading(text, level=level)
    heading.paragraph_format.page_break_before = True
    return heading


def add_scope_table(document: Document) -> None:
    document.add_heading("Completed Week 4 scope / 第四周已完成范围", level=1)
    table = document.add_table(rows=1, cols=3)
    widths = [3120, 3120, 3120]
    set_table_fixed(table, widths)
    headers = [
        "Core classification updates\n核心分类更新",
        "Necessary supporting work\n必要配套工作",
        "Pilot-scale enhancements\n试点范围内的增强",
    ]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        clear_cell(cell)
        set_cell_shading(cell, NAVY)
        set_cell_border(cell, color=WHITE, size="4")
        p = cell.paragraphs[0]
        add_run(p, header, bold=True, color=WHITE, size=9)
    repeat_header(table.rows[0])
    body = table.add_row()
    prevent_row_split(body)
    content = [
        [
            ("Separated Simulation and Case-Based Learning", "拆分模拟学习与案例式学习"),
            ("Applied non-exclusive multi-label classification", "采用非互斥的多标签分类"),
            ("Calibrated WIL weights", "校准 WIL 权重"),
            ("Added total scoring and configured all supplied categories", "增加总评分并配置全部既定类别"),
        ],
        [
            ("Reran INFS6600 only", "仅重新运行 INFS6600"),
            ("Added overlap and review policy", "增加类别重叠与人工复核政策"),
            ("Produced structured evidence and mapping outputs", "生成结构化证据与映射输出"),
            ("Added regression tests, release notes and audit hashes", "增加回归测试、发布说明与审计哈希"),
        ],
        [
            ("Counted distinct evidence items", "统计不同证据条目"),
            ("Created three INFS6600 category-distribution figures", "制作三张 INFS6600 分类分布图"),
            ("Kept both suggestions inside the existing pilot scope", "将两项建议严格限制在现有试点范围内"),
        ],
    ]
    for idx, lines in enumerate(content):
        cell = body.cells[idx]
        clear_cell(cell)
        set_cell_shading(cell, PALE_BLUE if idx < 2 else PALE_GOLD)
        set_cell_border(cell)
        for english, chinese in lines:
            add_bilingual_bullet(cell, english, chinese)


def add_allocation_overview(document: Document) -> None:
    add_page_heading(document, "Eight-person allocation at a glance / 八人分工总览")
    p = document.add_paragraph()
    add_run(p, "The table records the completed Week 4 workstream and meeting focus for each member.\n下表记录每位成员已经完成的第四周工作流及会议汇报重点。", italic=True, color=MUTED, size=8.7)

    table = document.add_table(rows=1, cols=4)
    widths = [600, 1830, 4650, 2280]
    set_table_fixed(table, widths)
    headers = ["#", "Member / 成员", "Primary workstream / 主要工作流", "Meeting focus / 汇报重点"]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        clear_cell(cell)
        set_cell_shading(cell, NAVY)
        set_cell_border(cell, color=WHITE, size="4")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx in (0, 3) else WD_ALIGN_PARAGRAPH.LEFT
        add_run(p, header, bold=True, color=WHITE, size=8.5)
    repeat_header(table.rows[0])

    focuses = [
        ("Scope and integration", "范围与整合"),
        ("Taxonomy changes", "分类体系修改"),
        ("Overlap policy", "重叠政策"),
        ("Evidence rerun", "证据重跑"),
        ("Scoring changes", "评分修改"),
        ("Results and mapping", "结果与映射"),
        ("Figures and reports", "图表与报告"),
        ("QA and release", "质量检查与发布"),
    ]
    for member, focus in zip(MEMBERS, focuses):
        row = table.add_row()
        prevent_row_split(row)
        fill = WHITE if int(member["number"]) % 2 else GREY
        values = (
            member["number"],
            member["name"],
            f"{member['stream']}\n{member['stream_zh']}",
            f"{focus[0]}\n{focus[1]}",
        )
        for idx, value in enumerate(values):
            cell = row.cells[idx]
            clear_cell(cell)
            set_cell_shading(cell, fill)
            set_cell_border(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx in (0, 3) else WD_ALIGN_PARAGRAPH.LEFT
            add_run(p, value, bold=idx == 1, color=NAVY if idx in (0, 1) else TEXT, size=7.8)

    document.add_heading("Team completion record / 团队完成记录", level=2)
    add_callout(
        document,
        "The eight-member CS-44 team completed all Week 4 workstreams, including taxonomy configuration, INFS6600 evidence analysis, scoring, visualisation, quality assurance and final delivery.\nCS-44 八人团队完成了第四周的全部工作流，包括分类体系配置、INFS6600 证据分析、评分、可视化、质量检查和最终交付。",
    )


def add_member_card(document: Document, member: dict) -> None:
    table = document.add_table(rows=1, cols=2)
    widths = [4680, 4680]
    set_table_fixed(table, widths)
    header = table.rows[0]
    merged = header.cells[0].merge(header.cells[1])
    clear_cell(merged)
    set_cell_shading(merged, BLUE_DARK)
    set_cell_border(merged, color=BLUE_DARK, size="8")
    p = merged.paragraphs[0]
    add_run(p, f"{member['number']}  {member['name']}", bold=True, color=WHITE, size=11)
    add_run(p, f"  |  {member['stream']} / {member['stream_zh']}", color=WHITE, size=8.8)

    body = table.add_row()
    prevent_row_split(body)
    for idx, heading in enumerate(("Completed responsibilities / 已完成职责", "Concrete deliverables / 具体交付物")):
        cell = body.cells[idx]
        clear_cell(cell)
        set_cell_shading(cell, WHITE if idx == 0 else PALE_BLUE)
        set_cell_border(cell)
        p = cell.paragraphs[0]
        add_run(p, heading, bold=True, color=BLUE, size=8.8)
        items = member["completed"] if idx == 0 else member["outputs"]
        items_zh = member["completed_zh"] if idx == 0 else member["outputs_zh"]
        for english, chinese in zip(items, items_zh):
            add_bilingual_bullet(cell, english, chinese)

    line_row = table.add_row()
    prevent_row_split(line_row)
    line_cell = line_row.cells[0].merge(line_row.cells[1])
    clear_cell(line_cell)
    set_cell_shading(line_cell, PALE_GOLD)
    set_cell_border(line_cell, color=GOLD, size="6")
    p = line_cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_run(p, "Suggested meeting line / 建议汇报话术:\n", bold=True, color=NAVY, size=8.4)
    add_run(p, f"“{member['line']}”\n", italic=True, color=TEXT, size=8.25)
    add_run(p, f"“{member['line_zh']}”", color=MUTED, size=8.25)

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_member_pages(document: Document) -> None:
    for start in range(0, len(MEMBERS), 2):
        end = start + 2
        add_page_heading(
            document,
            f"Completed contributions — Members {start + 1}–{end} / 成员 {start + 1}–{end} 已完成工作",
        )
        for member in MEMBERS[start:end]:
            add_member_card(document, member)


def add_final_page(document: Document) -> None:
    add_page_heading(document, "Meeting run sheet and shared wording / 会议流程与统一口径")
    add_callout(
        document,
        "Recommended total speaking time: 8–10 minutes. Each member owns one 40–60 second update; Member 1 opens with scope and Member 8 closes with verification and release status.\n建议总汇报时间为 8–10 分钟。每位成员负责 40–60 秒；Member 1 以范围说明开场，Member 8 以验证与发布状态收尾。",
    )

    document.add_heading("Speaking order / 汇报顺序", level=2)
    table = document.add_table(rows=1, cols=3)
    widths = [900, 2550, 5910]
    set_table_fixed(table, widths)
    for idx, header in enumerate(("Order / 顺序", "Speaker / 汇报人", "One-sentence hand-off / 一句话衔接")):
        cell = table.rows[0].cells[idx]
        clear_cell(cell)
        set_cell_shading(cell, NAVY)
        set_cell_border(cell, color=WHITE, size="4")
        p = cell.paragraphs[0]
        add_run(p, header, bold=True, color=WHITE, size=8.5)
    repeat_header(table.rows[0])
    handoffs = [
        ("Week 4 decisions, acceptance checklist and scope", "第四周决策、验收清单与范围"),
        ("Eight-category taxonomy and Simulation/Case split", "八分类体系及模拟/案例拆分"),
        ("Multi-label, overlap and review treatment", "多标签、重叠关系与复核处理"),
        ("INFS6600 evidence preparation and rerun", "INFS6600 证据准备与重跑"),
        ("Weights, scoring and provisional thresholds", "权重、评分与临时阈值"),
        ("Distinct counts, totals and category mapping", "不同证据计数、总分与类别映射"),
        ("Pilot-scale figures and reports", "试点规模图表与报告"),
        ("Tests, release manifest, PDFs and versioned Git delivery", "测试、发布清单、PDF 与版本化 Git 交付"),
    ]
    for idx, (member, handoff) in enumerate(zip(MEMBERS, handoffs), start=1):
        row = table.add_row()
        prevent_row_split(row)
        fill = WHITE if idx % 2 else GREY
        for col, value in enumerate((str(idx), member["name"], f"{handoff[0]}\n{handoff[1]}")):
            cell = row.cells[col]
            clear_cell(cell)
            set_cell_shading(cell, fill)
            set_cell_border(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col == 0 else WD_ALIGN_PARAGRAPH.LEFT
            add_run(p, value, bold=col == 1, color=NAVY if col == 1 else TEXT, size=8.1)

    document.add_heading("Approved team delivery statement / 统一团队交付口径", level=2)
    add_callout(
        document,
        "The eight-member CS-44 project team completed the Week 4 v2 release. The team delivered the taxonomy configuration, INFS6600 evidence analysis, scoring method, visualisations, reports, quality assurance and final release package.\nCS-44 八人项目团队完成了第四周 v2 版本。团队完成了分类体系配置、INFS6600 证据分析、评分方法、可视化、报告、质量检查及最终发布包。",
        fill=BLUE_LIGHT,
    )

    document.add_heading("Closing result statement / 总结口径", level=2)
    p = document.add_paragraph()
    add_run(
        p,
        "Together, the team implemented the agreed changes in the existing INFS6600 pilot, preserved multi-label classification, and confirmed positive allocation to Work-Integrated and Applied Learning, Case-Based Learning, and Project- and Problem-Based Learning. No additional units or unrelated future-stage work were advanced this week.\n团队共同在现有 INFS6600 试点中实施已确认修改，保留多标签分类，并确认课程可正向归入工作整合与应用学习、案例式学习以及项目与问题式学习。本周未推进其他课程或无关的后续阶段工作。",
        size=9.4,
    )

    document.add_heading("Scope guard for questions / 提问时的范围口径", level=2)
    guard = document.add_table(rows=1, cols=2)
    set_table_fixed(guard, [4680, 4680])
    for idx, (title, items, fill) in enumerate(
        (
            (
                "Completed this week / 本周已完成",
                [
                    ("INFS6600 v2 classification changes", "INFS6600 v2 分类修改"),
                    ("Necessary supporting implementation", "必要配套实施"),
                    ("Two pilot-scale enhancements", "两项试点范围内的增强"),
                ],
                PALE_BLUE,
            ),
            (
                "Not advanced this week / 本周未推进",
                [
                    ("Additional units or UG/PG comparison", "其他课程或本科/研究生比较"),
                    ("Landing page or LLM/RAG work", "落地页或 LLM/RAG 工作"),
                    ("Formal precision/recall/F1 evaluation", "正式的精确率/召回率/F1 评估"),
                ],
                GREY,
            ),
        )
    ):
        cell = guard.cell(0, idx)
        clear_cell(cell)
        set_cell_shading(cell, fill)
        set_cell_border(cell)
        p = cell.paragraphs[0]
        add_run(p, title, bold=True, color=NAVY, size=9)
        for english, chinese in items:
            add_bilingual_bullet(cell, english, chinese)


def build_document(path: Path) -> None:
    document = Document()
    configure_document(document)
    document.core_properties.title = "CS-44 Week 4 Eight-Person Work Allocation / 第四周八人分工"
    document.core_properties.subject = "Week 4 v2 bilingual meeting reporting notes / 第四周 v2 双语会议汇报材料"
    document.core_properties.author = "CS-44 Project Team (8 members)"
    document.core_properties.comments = "Prepared collaboratively for the Week 4 project meeting."

    p = document.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(p, "Week 4 Eight-Person Work Allocation", bold=True, color=NAVY, size=24)
    add_run(p, "\n第四周八人分工", bold=True, color=NAVY, size=21)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    add_run(subtitle, "INFS6600 Taxonomy Pilot - Week 4 v2 / INFS6600 分类体系试点 - 第四周 v2", bold=True, color=BLUE, size=11.6)
    add_run(subtitle, "\nMeeting reporting notes and completed contribution record / 会议汇报材料与已完成贡献记录", color=MUTED, size=9.6)

    add_metadata(document)
    add_callout(
        document,
        "The eight-member CS-44 team completed all Week 4 work and deliverables. The allocation below records the work completed by each member for the meeting presentation.\nCS-44 八人团队完成了第四周的全部工作和交付物。以下分工记录每位成员为会议汇报完成的具体工作。",
    )

    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    add_run(p, "Week 3 continuity basis / 与第三周工作的连续性", bold=True, color=NAVY, size=10)
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    add_run(
        p,
        "This Week 4 allocation continues the Week 3 team-level workflow from scope and taxonomy through extraction, classification, visualisation and handover. Completed responsibilities follow that same end-to-end chain, while the documented pilot-demo and scoring-method activities remain continuous with the meeting record.\n本周分工延续第三周从范围与分类体系、数据提取、分类分析、可视化到交付的团队工作流。各成员已完成职责按照同一端到端链条安排，并保持试点演示与评分方法等会议记录中已明确活动的连续性。",
        size=8.9,
    )

    add_scope_table(document)
    add_allocation_overview(document)
    add_member_pages(document)
    add_final_page(document)

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--output", type=Path, default=OUTPUT)
    args = parser.parse_args()
    build_document(args.output)
    print(f"Wrote {args.output}")


if __name__ == "__main__":
    main()
